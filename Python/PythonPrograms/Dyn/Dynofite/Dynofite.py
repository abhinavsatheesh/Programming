from tkinter import *
from tkinter import filedialog as fd
import docx
from tkinter import messagebox
import os
import sys
import docx2txt

#Version = Version 6       
print("Dynofite\nVersion 6")

def DYNOFITE():
    class DOCUMENT(Tk):   
        def __init__(self):
            super().__init__()
            self.title("Dynofite - The Only Office You Need | Word Document")
            self.state("zoomed")
            self.WELCOME=Label(self, text="Welcome to Dynofite - The Only Office You Need | Word Document")
            self.WELCOME.place(x=500, y=0)
            self.NAME=Label(self, text="Enter the name of the document").place(x=20, y=50)
            self.NAMEENTRY=Entry(self)
            self.NAMEENTRY.place(x=500, y=50)
            self.HEAD=Label(self, text="Enter the heading of the document")
            self.HEAD.place(x=20, y=90)
            self.e1=Entry(self)
            self.e1.place(x=500, y=90)
            self.PARA=Label(self, text="Enter the text or paragraph which should be included in the document")
            self.PARA.place(x=20, y=130)
            self.e2=Text(self)
            self.e2.place(x=500, y=130, height=200, width=750)
            self.importbutton = Button(self, text="Import a Word Document", command=self.IMPORTWORD).place(x=400, y=500)
            self.CREATEW=Button(self, text="Export as a Word Document", command=self.CREATEWORD)
            self.CREATEW.place(x=600, y=500)
            self.mainloop()
        
        def IMPORTWORD(self):
            PATH = fd.askopenfilename()
            my_text = docx2txt.process(PATH)
            self.e2.delete(1.0,"end")
            self.e2.insert(1.0, my_text)
            
        def CREATEWORD(self):
            self.doc = docx.Document()
            nameTEXT=self.NAMEENTRY.get()
            headTEXT=self.e1.get()
            paraTEXT=self.e2.get("1.0",'end-1c')
            if headTEXT=="":
                pass
            elif headTEXT!="":
                self.doc.add_heading(headTEXT, 0)
            if nameTEXT == "":
                messagebox.showinfo("No Name", "No name has been entered for the document. Please enter a name in the Name Column")
            else:
                try:
                    filename1 = fd.askdirectory()
                    print(filename1)
                    if paraTEXT=="":
                        directory=f"{filename1}/{nameTEXT}.docx"
                        self.doc.save(directory)
                        os.startfile(directory)
                        exit()
                    else:
                        self.doc.add_paragraph(paraTEXT)
                        directory=f"{filename1}/{nameTEXT}.docx"
                        self.doc.save(directory)
                        os.startfile(directory)
                        exit()
                except:
                    pass
                
                   
    documentApp = DOCUMENT()
DYNOFITE()
